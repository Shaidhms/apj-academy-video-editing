"""Build the Word document for the CapCut step-by-step video editing guide."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/shaid/Documents/Shaid_Works/Nisha/video-editing-guide/CapCut-Video-Editing-Guide.docx"

doc = Document()

# ---- Styles ----
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def H1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)


def H2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)


def H3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x3D, 0x6E, 0x99)


def P(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def BULLET(text):
    doc.add_paragraph(text, style="List Bullet")


def NUMBER(text):
    doc.add_paragraph(text, style="List Number")


def NOTE(text):
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB0, 0x55, 0x00)
    run2 = p.add_run(text)
    run2.italic = True


def TIP(text):
    p = doc.add_paragraph()
    run = p.add_run("Tip: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x6E, 0x4F)
    p.add_run(text)


def WARN(text):
    p = doc.add_paragraph()
    run = p.add_run("Warning: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB0, 0x1A, 0x1A)
    p.add_run(text)


# ===== TITLE PAGE =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("CapCut Video Editing")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Complete Step-by-Step Guide for Content Creators")
r.font.size = Pt(14)
r.italic = True
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Prepared for: Shaid\n").bold = True
meta.add_run("Topics covered: Installation (with VPN), Editing Basics, "
             "Keyframes, Transitions, Effects, Text Masking, "
             "Color Grading & HSL, and the CapCut Pro Export Trick").italic = True

doc.add_paragraph()
doc.add_paragraph(
    "This guide is a draft for verification. Once you confirm the content, "
    "the same document will be exported as a PDF."
).italic = True

doc.add_page_break()

# ===== TABLE OF CONTENTS =====
H1("Table of Contents")
toc = [
    "1. Before You Start: Tools You Will Need",
    "2. Installing CapCut in India Using a VPN (Step-by-Step)",
    "3. CapCut Interface Tour (PC)",
    "4. Splitting and Cutting Clips",
    "5. Keyframes: How to Animate Anything",
    "6. Transitions Between Clips",
    "7. Effects (Video, Body, and Overlay Effects)",
    "8. Text Masking: Reveal Text Effect",
    "9. Color Grading Basics (Brightness, Contrast, Temperature, Tint)",
    "10. HSL Adjustment for Cinematic Color",
    "11. Color Gradations and Pro Looks",
    "12. The CapCut Pro Export Trick (No Watermark, Free)",
    "13. Export Settings Cheat Sheet",
    "14. Common Mistakes to Avoid",
    "15. References & Sources",
]
for item in toc:
    doc.add_paragraph(item)

doc.add_page_break()

# ===== SECTION 1 =====
H1("1. Before You Start: Tools You Will Need")
P("To follow this guide end to end, keep these ready before opening CapCut:")
BULLET("A Windows PC or Mac (this guide focuses on the PC version, "
       "since the Pro export trick works only on PC).")
BULLET("A stable internet connection. CapCut downloads assets and effects "
       "on demand the first time you use them.")
BULLET("A VPN app (free or paid). Recommended: Proton VPN (free tier works), "
       "Urban VPN, Hotspot Shield, or NordVPN if you want a paid option.")
BULLET("A Google or TikTok account to sign into CapCut.")
BULLET("Around 4 GB of free disk space for the app, cache and exports.")
BULLET("Your raw video clips, music, and any logo or overlay PNGs ready in "
       "one folder so you can drag them in quickly.")

NOTE("CapCut is officially banned in India since 2020. The app is not "
     "available on the Indian Play Store, App Store, or Microsoft Store. "
     "Individual users have not been prosecuted for using it, but this is "
     "why you need a VPN to install it.")

# ===== SECTION 2 =====
H1("2. Installing CapCut in India Using a VPN (Step-by-Step)")

H2("Method A — On a Windows PC")
P("This is the cleanest method and is the same one used in most YouTube "
  "tutorials, including the Tamil tutorial you referenced.")

H3("Step 1: Change your Windows region")
NUMBER("Click the Start button and open Settings (or press Windows + I).")
NUMBER("Go to Time & Language.")
NUMBER("Click Language & Region.")
NUMBER("Under Country or region, change it from India to United Kingdom "
       "or United States.")
NUMBER("Close Settings.")

NOTE("This step alone unlocks CapCut in the Microsoft Store. You only need "
     "to do this once.")

H3("Step 2: Install a free VPN")
NUMBER("Open your browser and download Hotspot Shield, Proton VPN, "
       "or Urban VPN from their official website. Avoid random APK or "
       "EXE links from forums.")
NUMBER("Run the installer and follow the on-screen prompts.")
NUMBER("Open the VPN app and sign up for the free plan.")
NUMBER("Connect to a server in the United Kingdom or United States.")
NUMBER("Wait until the VPN shows Connected before moving to the next step.")

WARN("Do not skip the VPN. If you open CapCut without a VPN active, "
     "you will see a Network Error and the app will not load even after "
     "you change the region.")

H3("Step 3: Download CapCut")
NUMBER("Go to https://www.capcut.com in your browser (with VPN still on).")
NUMBER("Click the Download for Windows button on the home page.")
NUMBER("Once the .exe file finishes downloading, double-click it.")
NUMBER("Follow the installer. Default options are fine. Choose your install "
       "location if you want it on a non-system drive.")
NUMBER("Wait until installation completes and click Finish.")

TIP("Always download the standalone installer from capcut.com, not the "
    "Microsoft Store version. The Pro export trick in Section 12 only "
    "works on the standalone installer.")

H3("Step 4: First-time launch and sign in")
NUMBER("Make sure your VPN is still connected to UK or USA.")
NUMBER("Launch CapCut from the desktop shortcut.")
NUMBER("Click Sign in and choose Google, TikTok, or Facebook.")
NUMBER("Complete the sign-in. You should now see the CapCut home screen "
       "with New Project and recent projects.")

H3("Step 5: After installation")
NUMBER("You can disconnect the VPN now for normal editing.")
NUMBER("Reconnect the VPN only when you want to sign in again or update "
       "the app.")
NUMBER("CapCut works offline for most editing. You only need internet for "
       "downloading new effects, transitions and stock assets.")

H2("Method B — On Android (Quick Reference)")
NUMBER("Install Proton VPN from the Play Store and connect to a UK or US server.")
NUMBER("Open Play Store, go to Settings → General → Account preferences → "
       "Country and profiles. Wait for the country to update to UK or US.")
NUMBER("Clear Play Store cache (Settings → Apps → Play Store → Storage → "
       "Clear cache).")
NUMBER("Search CapCut on Play Store. It should now appear. Install it.")
NUMBER("Open CapCut with VPN still on for the first launch and sign in.")
NUMBER("After first launch, the app works without a VPN for most editing.")

WARN("Avoid downloading CapCut Pro APK files from random websites. "
     "Many of them carry malware. Stick to the official Play Store or "
     "App Store after the region change.")

doc.add_page_break()

# ===== SECTION 3 =====
H1("3. CapCut Interface Tour (PC)")
P("Before you edit, get familiar with the layout. CapCut PC has 4 main areas:")
BULLET("Top-left: Media library, Audio, Text, Stickers, Effects, "
       "Transitions, Filters, Adjust.")
BULLET("Top-right: Preview window where you watch your edit.")
BULLET("Bottom: Timeline where you arrange clips, audio and text.")
BULLET("Right side: Properties panel where you edit the selected clip "
       "(scale, position, opacity, color, animation, etc.).")

P("Useful shortcuts you will use constantly:")
BULLET("Ctrl + B — Split clip at the playhead.")
BULLET("Ctrl + I — Import media.")
BULLET("Ctrl + Z / Ctrl + Y — Undo / Redo.")
BULLET("Spacebar — Play / Pause preview.")
BULLET("Delete — Remove selected clip or segment.")
BULLET("Alt + K — Open the keyframe panel.")

# ===== SECTION 4 =====
H1("4. Splitting and Cutting Clips")
P("Splitting is the foundation of every edit. You cut a long clip into "
  "smaller pieces, throw away the parts you do not want, and rearrange "
  "the rest.")

H2("Step-by-Step: Split a Clip")
NUMBER("Open CapCut and click New Project.")
NUMBER("Click Import (or press Ctrl + I) and add your video.")
NUMBER("Drag the video from the media library down to the timeline.")
NUMBER("Use the spacebar to play the video and find the exact moment you "
       "want to cut.")
NUMBER("Pause when the playhead (the vertical white line) is on that moment.")
NUMBER("Click the scissors / Split icon on the toolbar above the timeline. "
       "Or press Ctrl + B.")
NUMBER("The clip is now cut into two separate pieces at that exact frame.")

H2("Step-by-Step: Remove an Unwanted Section")
NUMBER("Move the playhead to the start of the part you want to delete.")
NUMBER("Press Ctrl + B to split.")
NUMBER("Move the playhead to the end of the part you want to delete.")
NUMBER("Press Ctrl + B to split again.")
NUMBER("Click on the middle segment (the unwanted part) to select it.")
NUMBER("Press Delete on your keyboard.")
NUMBER("The remaining clips automatically snap together. Done.")

TIP("Cut on action and on the beat. If your video has music, split where "
    "the beat hits. If there is movement (a hand reaching, a head turning) "
    "cut just as the movement starts. The result feels professional even "
    "if you do nothing else.")

# ===== SECTION 5 =====
H1("5. Keyframes: How to Animate Anything")
P("A keyframe is a marker that tells CapCut: at this exact moment, this "
  "property has this value. When you set two keyframes with different "
  "values, CapCut animates smoothly between them.")

P("You can keyframe almost anything: position, scale, rotation, opacity, "
  "volume, blur, color values, and more.")

H2("Step-by-Step: Create a Zoom-In Keyframe Animation")
NUMBER("Drag a clip to the timeline.")
NUMBER("Click the clip once to select it. The right-side properties panel "
       "appears.")
NUMBER("Move the playhead to the start of the clip.")
NUMBER("In the right panel, find Scale (under Basic). Note the diamond "
       "icon next to it.")
NUMBER("Click that diamond icon to add the first keyframe at the current "
       "position. The diamond turns orange.")
NUMBER("Now move the playhead about 2 seconds forward.")
NUMBER("Change the Scale value to 120 percent. CapCut automatically adds "
       "a second keyframe.")
NUMBER("Press the spacebar to preview. The clip slowly zooms from 100 to "
       "120 percent.")

H2("Step-by-Step: Create a Fade-In with Keyframes")
NUMBER("Select your clip on the timeline.")
NUMBER("Move the playhead to the start of the clip.")
NUMBER("In the right panel, set Opacity to 0 percent and click the diamond "
       "to add a keyframe.")
NUMBER("Move the playhead 1 second forward.")
NUMBER("Change Opacity to 100 percent. A second keyframe is added.")
NUMBER("Play. The clip fades in over 1 second.")

H2("Useful Keyframe Tricks")
BULLET("Position keyframes — make text fly in from the left, right, or "
       "bottom of the screen.")
BULLET("Rotation keyframes — rotate a logo 360 degrees over 2 seconds for "
       "a stinger effect.")
BULLET("Scale keyframes — push in on a face during emotional moments.")
BULLET("Volume keyframes (audio) — duck music down when someone speaks "
       "and ramp it up between sentences.")

TIP("To delete a keyframe, right-click on it in the timeline panel "
    "(or in the keyframe row) and choose Delete.")

# ===== SECTION 6 =====
H1("6. Transitions Between Clips")
P("Transitions are the visual effects that play between two clips: fades, "
  "wipes, blurs, zooms, glitches, and so on. Used well, they make your "
  "edit feel intentional. Used too much, they look amateur.")

H2("Step-by-Step: Add a Transition")
NUMBER("Place two clips next to each other on the timeline so they touch.")
NUMBER("Click the Transitions tab in the top-left media panel.")
NUMBER("Browse the categories: Basic, Camera, Light Effect, Glitch, "
       "Distortion, Slide, Mask, MG, Social Media, Overlay, Blur.")
NUMBER("Hover over a transition to preview it.")
NUMBER("If a download icon appears, click it once. CapCut will fetch the "
       "transition asset.")
NUMBER("Drag the transition between the two clips on the timeline. A "
       "small box appears at the cut point.")
NUMBER("Click the transition box. In the right panel, change Duration to "
       "set how long it lasts (e.g. 0.5 seconds).")

NOTE("The maximum duration of a transition is half the length of the "
     "shorter neighboring clip. If one clip is 4 seconds and the other "
     "is 3 seconds, the transition can be up to 1.5 seconds.")

H2("Step-by-Step: Apply One Transition to All Clips")
NUMBER("Select the transition you want.")
NUMBER("Drag it between the first two clips.")
NUMBER("Right-click the transition box on the timeline.")
NUMBER("Choose Apply to all.")
NUMBER("Every cut on that track now uses the same transition.")

TIP("For talking-head and vlog content, stick to soft transitions like "
    "Dissolve, Whip Pan, or Zoom Blur. Save Glitch and Distortion for "
    "music or hype edits.")

# ===== SECTION 7 =====
H1("7. Effects: Video, Body, and Overlay Effects")
P("Effects in CapCut are filters or animated graphics that sit on top of "
  "your clip. They are organized into Video Effects (apply to the whole "
  "frame, like film grain or VHS) and Body Effects (track a person, like "
  "a glow outline).")

H2("Step-by-Step: Add a Video Effect")
NUMBER("Click the Effects tab in the top-left panel.")
NUMBER("Choose Video Effects.")
NUMBER("Browse categories — Trending, Basic, Bling, Vintage, Atmosphere, "
       "Fantasy, etc.")
NUMBER("Click an effect to preview. If a download icon shows, click it.")
NUMBER("Drag the effect onto the timeline above your clip. It will land "
       "on its own track.")
NUMBER("Stretch or shrink the effect bar to control how long the effect "
       "plays.")
NUMBER("Click the effect on the timeline. In the right panel, adjust "
       "intensity, speed, and color sliders.")

H2("Step-by-Step: Add an Overlay (Picture-in-Picture or Stock Asset)")
NUMBER("Press Ctrl + I and import the overlay file (a video, PNG, or "
       "stock clip).")
NUMBER("Drag the overlay above your main clip on the timeline.")
NUMBER("In the preview window, drag the corner handles to resize the "
       "overlay and reposition it.")
NUMBER("In the right panel, find Blend mode and try Screen, Add, Overlay, "
       "or Multiply for cinematic effects (great for light leaks, "
       "smoke, dust, or rain overlays).")
NUMBER("Lower Opacity if the overlay is too strong.")

TIP("For light leaks and dust overlays from YouTube, set blend mode to "
    "Screen. For dark overlays like rain on glass, use Multiply.")

# ===== SECTION 8 =====
H1("8. Text Masking: Reveal Text Effect")
P("Text masking lets you put video inside text, or make text appear "
  "naturally in a scene (like behind a person walking past). It is one of "
  "the most reused effects on Reels and TikTok.")

H2("Step-by-Step: Video Inside Text (Beginner)")
NUMBER("Drag your video to the timeline.")
NUMBER("Click Text → Add default text. A text box appears in the preview.")
NUMBER("Type a short word like BOLD or your name. Use a thick, fat font "
       "(Impact, Anton, Bebas) — thin fonts do not work for masking.")
NUMBER("Make the font huge, around 200 pt or larger. Position it where "
       "you want.")
NUMBER("Stretch the text bar on the timeline so it covers the entire clip "
       "duration.")
NUMBER("Right-click the text and choose Convert to overlay (or duplicate "
       "the video and place it above the text).")
NUMBER("Apply a Mask shaped like the text using the text layer as the "
       "mask source.")
NUMBER("Play. Your video now plays only inside the letters.")

H2("Step-by-Step: Walk-Behind Text Reveal (Advanced)")
P("This is the trick where text appears like it is behind a person.")

NUMBER("Drag your clip to the main timeline (Track 1).")
NUMBER("Duplicate the clip onto Track 3 above it (right-click → Copy, "
       "then paste above).")
NUMBER("Add your text on Track 2 (between the two video layers).")
NUMBER("Select the top video layer (Track 3).")
NUMBER("In the right panel, click Mask → choose the Rectangle or Custom "
       "shape.")
NUMBER("Position the mask so only the part of the body that should cover "
       "the text is visible (e.g. just the person's silhouette).")
NUMBER("Use keyframes (see Section 5) on the mask Position so the mask "
       "follows the person frame by frame as they walk past the text.")
NUMBER("Play it back. The text appears to be behind the person.")

TIP("Filming tip — shoot with a static camera (tripod) when you plan to "
    "use this effect. Camera movement makes masking 10x harder.")

# ===== SECTION 9 =====
H1("9. Color Grading Basics")
P("Color grading is the process of changing the look and mood of your "
  "footage with color. Before you touch HSL, get the basics right.")

H2("Step-by-Step: Basic Color Correction")
NUMBER("Select your clip on the timeline.")
NUMBER("In the right panel, click Adjust.")
NUMBER("Work through the sliders top to bottom:")
BULLET("Brightness — pull up if the shot is dark, down if blown out.")
BULLET("Contrast — adds punch. Most phone footage is too flat; +10 to +20 "
       "is a good starting point.")
BULLET("Saturation — too high looks fake. Stay between -10 and +15.")
BULLET("Sharpness — only push if the footage is soft; +5 to +10 max.")
BULLET("Highlights — pull down to recover sky details.")
BULLET("Shadows — pull up to see detail in dark areas, but do not "
       "destroy mood.")
BULLET("Temperature — warm (right) for golden-hour feel, cool (left) for "
       "moody/cinematic.")
BULLET("Tint — push slightly toward green or magenta to fix color casts.")

TIP("Always grade on a calibrated screen. Phone screens are often too "
    "saturated, so what looks great on your phone can look flat on a "
    "laptop. If possible, view the export on the platform you publish on "
    "(Instagram, YouTube) before sharing.")

# ===== SECTION 10 =====
H1("10. HSL Adjustment for Cinematic Color")
P("HSL stands for Hue, Saturation, and Lightness. Unlike the global "
  "saturation slider, HSL lets you change one specific color without "
  "touching the others. This is how every cinematic edit gets the teal "
  "and orange look.")

H2("Step-by-Step: Use the HSL Panel")
NUMBER("Select your clip.")
NUMBER("Click Adjust in the right panel.")
NUMBER("Scroll down and click HSL.")
NUMBER("You will see 8 color circles: Red, Orange, Yellow, Green, Aqua, "
       "Blue, Purple, Pink.")
NUMBER("Click on a color circle (e.g. Red).")
NUMBER("Three sliders appear:")
BULLET("Hue — shifts the color (e.g. push Reds toward Orange to warm up "
       "skin tones).")
BULLET("Saturation — how strong that color appears.")
BULLET("Lightness — how bright that color appears.")
NUMBER("Repeat for each color you want to control.")

H2("Recipe: Cinematic Teal & Orange Look")
NUMBER("Open HSL.")
NUMBER("Click Orange. Push Saturation +15 and Lightness +5 to pop "
       "skin tones.")
NUMBER("Click Yellow. Pull Saturation -20 and shift Hue slightly toward "
       "Orange so yellow walls do not look sickly.")
NUMBER("Click Aqua. Pull Saturation +10.")
NUMBER("Click Blue. Push Saturation +20, pull Lightness -10. This deepens "
       "skies and shadows.")
NUMBER("Click Green. Pull Saturation -15 and shift Hue toward Aqua so "
       "trees and grass do not look radioactive.")
NUMBER("Compare with the original by toggling Adjust off and on. Tweak "
       "until it feels natural, not extreme.")

H2("Recipe: Moody / Filmic Look")
NUMBER("Pull Highlights down and Shadows up slightly to compress contrast.")
NUMBER("Lower overall Saturation by 10 to 15.")
NUMBER("In HSL, lower Lightness on Reds and Oranges by -10 so skin is "
       "less rosy.")
NUMBER("Push Blues toward Cyan and lower Lightness on Blues by -10.")
NUMBER("Add a slight green tint by nudging Tint to the left.")

TIP("Save your HSL recipe as a preset. In the Adjust panel, click the "
    "Save as preset icon. Now you can apply the same look to every clip "
    "in one click.")

# ===== SECTION 11 =====
H1("11. Color Gradations and Pro Looks")
P("Once you understand basic color and HSL, gradations are about layering "
  "looks for a specific mood.")

H2("Layering: How Pros Build a Look")
NUMBER("Step 1 — Correction: fix exposure, white balance, and contrast first.")
NUMBER("Step 2 — Base look: apply a Filter or LUT (CapCut has dozens "
       "under Filters → Movie or Vlogger). Lower its Intensity to around "
       "40 to 60 percent so it does not overpower.")
NUMBER("Step 3 — HSL refinement: target individual colors as in Section 10.")
NUMBER("Step 4 — Final polish: subtle vignette, grain, and tiny "
       "Sharpness boost.")

H2("Common Color Grades and When to Use Them")
BULLET("Warm Golden — outdoor lifestyle, family content, sunset shots. "
       "Push Temperature +10, Highlights -5, Saturation +10.")
BULLET("Teal & Orange — talking head, cinematic shorts, travel. Use "
       "the HSL recipe above.")
BULLET("High Contrast B&W — emotional or moody storytelling. Saturation "
       "-100, Contrast +20, Shadows -10.")
BULLET("Pastel / Soft — beauty, fashion, and aesthetic Reels. Pull "
       "Saturation -15, Highlights -10, Lightness on skin tones +10.")
BULLET("Faded Film — vintage / nostalgic. Use the Faded preset, lift "
       "Shadows up to 20, lower Contrast by 10, add film grain effect "
       "at low intensity.")

# ===== SECTION 12 =====
H1("12. The CapCut Pro Export Trick (Free, No Watermark)")
P("This is the trick from the YouTube video at "
  "https://youtu.be/Zm4i-YT5M1k. It uses CapCut's own "
  "pre-processing engine to render Pro effects into a hidden cache "
  "folder, which you then re-import. The trick works on the standalone "
  "Windows installer (capcut.com), not on the Microsoft Store version.")

WARN("This is a workaround, not a hack. Use it for personal learning. "
     "If you go commercial, buy CapCut Pro to support the creators who "
     "build this software.")

H2("Step 1: Make Sure You Are on the Right Version")
NUMBER("Open CapCut. In the bottom-left, click your profile and view "
       "the version number.")
NUMBER("The trick works most reliably on CapCut PC version 5.x. Newer "
       "versions (7.x and above) sometimes encrypt the cache files and "
       "the trick may fail.")
NUMBER("If you are on a newer version and the trick fails, uninstall "
       "and download an older standalone installer from a trusted source.")

H2("Step 2: Build Your Edit With Pro Effects")
NUMBER("Edit your video normally — cuts, transitions, text, music.")
NUMBER("Apply any Pro effects, filters, or AI features you want. They "
       "will show the Pro tag in the editor — that is fine.")
NUMBER("Preview the full edit and confirm everything looks right.")

H2("Step 3: Convert Everything to a Compound Clip")
NUMBER("Press Ctrl + A on the timeline to select all clips, audio, "
       "text, and effects.")
NUMBER("Right-click the selection.")
NUMBER("Choose Create compound clip.")
NUMBER("Your timeline now shows one merged green block.")

H2("Step 4: Pre-process the Compound Clip")
NUMBER("Right-click the new compound clip.")
NUMBER("Choose Pre-process compound clip (sometimes labelled as "
       "Pre-render).")
NUMBER("A progress bar appears. Wait until it reaches 100 percent. This "
       "renders the full video with Pro effects into a cache folder, "
       "bypassing the export paywall check.")

H2("Step 5: Open the Hidden Cache Folder")
NUMBER("Press Windows key + R to open the Run dialog.")
NUMBER("Type appdata and press Enter.")
NUMBER("You will land in C:\\Users\\YourName\\AppData\\Roaming. Now "
       "click the address bar and replace Roaming with Local. Press Enter.")
NUMBER("Navigate to: Local → CapCut → User Data → Projects → "
       "com.lveditor.draft → [Your Project Name] → Resources → Combination.")
NUMBER("Inside Combination you will see one or more files ending in .mp4. "
       "Ignore any .alpha files.")
NUMBER("Sort by Date modified to find the latest one — this is the "
       "pre-processed video.")

NOTE("If you cannot see the AppData folder, enable hidden items in File "
     "Explorer: View → Show → Hidden items.")

H2("Step 6: Re-import and Export")
NUMBER("Copy that .mp4 to your Desktop or any folder you can find easily.")
NUMBER("Go back to CapCut.")
NUMBER("Press Ctrl + I and import the .mp4 you just copied.")
NUMBER("Drag the imported file onto the existing compound clip on the "
       "timeline. CapCut will ask if you want to Replace clip — choose Yes.")
NUMBER("Click Export in the top-right.")
NUMBER("Set resolution (1080p or 4K), frame rate (30 or 60 fps), and "
       "format (MP4).")
NUMBER("Click Export. The video saves with all Pro effects intact and "
       "no watermark.")

H2("If the Trick Fails")
BULLET("Make sure you used the standalone installer, not the Microsoft "
       "Store version.")
BULLET("Try downgrading CapCut to version 5.x.")
BULLET("Make sure pre-processing reached 100 percent before you went "
       "looking for the .mp4.")
BULLET("If the .mp4 plays back without effects, the cache was deleted; "
       "redo step 4 and retrieve immediately.")

# ===== SECTION 13 =====
H1("13. Export Settings Cheat Sheet")
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Platform"
hdr[1].text = "Recommended Settings"
hdr[2].text = "Notes"

rows = [
    ("Instagram Reels / TikTok", "1080x1920, 30 fps, MP4, H.264, 8-12 Mbps",
     "Vertical 9:16. Keep under 60s for max reach."),
    ("YouTube Shorts", "1080x1920, 30 or 60 fps, MP4, 10 Mbps",
     "Vertical 9:16. Up to 60s."),
    ("YouTube Long-form", "1920x1080 or 3840x2160, 60 fps, MP4, 25-50 Mbps",
     "Use 4K only if you shot in 4K."),
    ("LinkedIn", "1920x1080, 30 fps, MP4, 8 Mbps",
     "16:9 or 1:1 work well."),
    ("WhatsApp Status", "1080x1920, 30 fps, MP4, lower bitrate ~4 Mbps",
     "WhatsApp re-compresses; do not waste bitrate."),
]
for r in rows:
    row = table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

# ===== SECTION 14 =====
H1("14. Common Mistakes to Avoid")
BULLET("Over-using transitions. One or two styles per video, max.")
BULLET("Pushing Saturation and Sharpness too far. Subtle wins.")
BULLET("Forgetting to check the audio levels before exporting. Aim for "
       "voice peaks around -6 dB and music around -18 dB.")
BULLET("Using thin fonts for text masks. Always go thick and bold.")
BULLET("Forgetting to deselect a keyframe before changing a property "
       "globally. You may end up creating accidental keyframes.")
BULLET("Editing on a small phone screen and not checking on a laptop. "
       "Colors look different on bigger screens.")
BULLET("Skipping the VPN on first launch in India. CapCut will refuse "
       "to start.")

# ===== SECTION 15 =====
H1("15. References & Sources")
P("Reference video used for the Pro export trick:")
BULLET("How to use CapCut Pro for free in PC (Tamil) — "
       "https://youtu.be/Zm4i-YT5M1k")

P("Other references consulted while preparing this guide:")
BULLET("CapCut official guide on PC — "
       "https://www.capcut.com/resource/how-to-use-capcut-on-pc")
BULLET("Keyframes guide — "
       "https://www.capcut.com/resource/how-to-add-keyframes-in-capcut")
BULLET("HSL color grading — "
       "https://www.capcut.com/tools/hsl-color")
BULLET("Color grading masterclass — "
       "https://www.capcut.com/resource/capcut-color-grading")
BULLET("Transitions guide — "
       "https://www.tourboxtech.com/en/news/capcut-transitions.html")
BULLET("Text masking tutorial — "
       "https://www.youtube.com/watch?v=oJdVRVzLvVw")
BULLET("Using CapCut in India (VPN guide) — "
       "https://kripeshadwani.com/how-to-use-capcut-in-india/")
BULLET("CapCut Pro free trick (PC) — "
       "https://techydruid.com/get-capcut-pro-free-pc/")
BULLET("CapCut Pro free trick variations — "
       "https://www.ruhanirabin.com/unlock-capcut-pro-windows-free/")

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("End of draft. Verify the content, then this will be "
                 "exported as a PDF.")
r.italic = True
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.save(OUT)
print(f"Saved: {OUT}")
