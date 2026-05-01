"""Build v2 of the video editing guide with images and 4 tools.
Tools covered: VPN (Android), CapCut, Clipchamp, LumaFusion."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = "/Users/shaid/Documents/Shaid_Works/Nisha/video-editing-guide"
IMG = os.path.join(ROOT, "images")
OUT = os.path.join(ROOT, "Video-Editing-Master-Guide.docx")

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)


def H1(text, page_break_before=True):
    if page_break_before:
        doc.add_page_break()
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)


def H2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)


def H3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x3D, 0x6E, 0x99)


def P(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic


def BUL(text):
    doc.add_paragraph(text, style="List Bullet")


def NUM(text):
    doc.add_paragraph(text, style="List Number")


def NOTE(text):
    p = doc.add_paragraph()
    r = p.add_run("Note: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xB0, 0x55, 0x00)
    r2 = p.add_run(text)
    r2.italic = True


def TIP(text):
    p = doc.add_paragraph()
    r = p.add_run("Tip: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x0B, 0x6E, 0x4F)
    p.add_run(text)


def WARN(text):
    p = doc.add_paragraph()
    r = p.add_run("Warning: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xB0, 0x1A, 0x1A)
    p.add_run(text)


def IMG_ADD(filename, caption=None, width_inches=5.5):
    """Insert image with optional caption, centered."""
    path = os.path.join(IMG, filename)
    if not os.path.exists(path) or os.path.getsize(path) < 1000:
        # placeholder text if missing
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[Screenshot missing: {filename}]")
        r.italic = True
        r.font.color.rgb = RGBColor(0x99, 0x55, 0x00)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(f"Figure: {caption}")
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def YT(label, url):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(f"{label} — ")
    r.bold = True
    p.add_run(url)


# ============================================================
# COVER PAGE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Video Editing Master Guide")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("VPN Setup • CapCut • Clipchamp • LumaFusion")
r.font.size = Pt(16)
r.italic = True
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
intro.add_run(
    "A complete step-by-step reference covering installation, splitting, "
    "keyframes, transitions, effects, text masking, color grading, HSL, "
    "and the CapCut Pro export trick, with screenshots, YouTube tutorials "
    "and tool-by-tool comparisons."
).italic = True

doc.add_paragraph()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run("Prepared for: Shaid\n").bold = True
m.add_run("Format: Word draft for verification, then export to PDF").italic = True

doc.add_paragraph()
doc.add_paragraph()
warn_box = doc.add_paragraph()
warn_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
wr = warn_box.add_run(
    "All screenshots in this document are sourced from public tutorials "
    "(Microsoft Clipchamp Blog, Wondershare, Kripesh Adwani, Proton VPN, "
    "Podfeet, Creative Bloq, AlphR). They are reproduced here for "
    "personal learning reference only. Source URLs are listed under "
    "each section."
)
wr.italic = True
wr.font.size = Pt(9)
wr.font.color.rgb = RGBColor(0x88, 0x55, 0x00)

# ============================================================
# TOC
# ============================================================
H1("Table of Contents")
toc = [
    "Part 1 — VPN: Why You Need One and How to Set It Up",
    "  1.1 Why a VPN matters for video apps in India",
    "  1.2 Installing a VPN on Android (Proton VPN)",
    "  1.3 Installing a VPN on Windows PC",
    "  1.4 Installing a VPN on iPhone / iPad",
    "",
    "Part 2 — CapCut (PC + Android)",
    "  2.1 Installing CapCut in India",
    "  2.2 CapCut interface tour",
    "  2.3 Splitting and cutting clips",
    "  2.4 Keyframes for animation",
    "  2.5 Transitions between clips",
    "  2.6 Effects and overlays",
    "  2.7 Text masking",
    "  2.8 Color grading basics",
    "  2.9 HSL adjustment for cinematic color",
    "  2.10 Color gradations and pro looks",
    "  2.11 CapCut Pro export trick (no watermark, free)",
    "",
    "Part 3 — Microsoft Clipchamp (Web + Windows 11)",
    "  3.1 What Clipchamp is and how it differs from CapCut",
    "  3.2 Getting started",
    "  3.3 Splitting clips",
    "  3.4 Keyframes (and the workaround)",
    "  3.5 Transitions",
    "  3.6 Effects and filters",
    "  3.7 Text and titles",
    "  3.8 Color adjustments",
    "  3.9 Exporting",
    "",
    "Part 4 — LumaFusion (iPhone / iPad / Android)",
    "  4.1 Why LumaFusion is different",
    "  4.2 Getting LumaFusion",
    "  4.3 Interface and timeline",
    "  4.4 Splitting and trimming",
    "  4.5 Keyframes",
    "  4.6 Transitions",
    "  4.7 Effects",
    "  4.8 Color grading and LUTs",
    "  4.9 Exporting",
    "",
    "Part 5 — Tool Comparison Cheat Sheet",
    "Part 6 — YouTube Tutorial Library",
    "Part 7 — Common Mistakes to Avoid",
    "Part 8 — References & Sources",
]
for line in toc:
    if line.strip():
        doc.add_paragraph(line)
    else:
        doc.add_paragraph()

# ============================================================
# PART 1 — VPN
# ============================================================
H1("Part 1 — VPN: Why You Need One and How to Set It Up")

H2("1.1 Why a VPN matters for video apps in India")
P("CapCut is officially banned in India since 2020 along with several "
  "other Chinese-origin apps. The app does not appear on Indian Play Store "
  "or App Store searches and the desktop version refuses to run if it "
  "detects an Indian IP. A VPN solves this by routing your internet "
  "traffic through a server in another country (UK, USA, Singapore), "
  "making the app think you are abroad.")

P("You will need a VPN for three reasons:")
BUL("Initial install — the app stores hide CapCut in India until your "
    "Play Store / Microsoft Store region is changed AND a VPN is active.")
BUL("First launch — CapCut PC throws a Network Error on Indian IPs even "
    "after install.")
BUL("Some Pro features and stock content libraries are geo-blocked.")

NOTE("Individual users have not been prosecuted for using CapCut. The "
     "ban is on app distribution, not personal use. Still, prefer a "
     "reputable VPN over random free APKs to stay safe.")

P("Recommended free / freemium VPNs:", bold=True)
BUL("Proton VPN — open source, free tier, no logs, no ads (best choice).")
BUL("Windscribe — 10 GB free per month, multiple regions.")
BUL("Hotspot Shield — fast, popular, free tier with ads.")
BUL("Urban VPN — completely free, lightweight, browser extension also.")

H2("1.2 Installing a VPN on Android (Proton VPN)")

H3("Step 1: Download Proton VPN from the Play Store")
NUM("Open Google Play Store on your Android phone.")
NUM("Search for Proton VPN.")
NUM("Tap Install. Wait for it to finish.")
NUM("Tap Open after installation.")

H3("Step 2: Create a Proton account (free)")
NUM("Tap Create account on the welcome screen.")
NUM("Enter a username and a strong password. Email is optional but "
    "recommended for recovery.")
NUM("You will be auto-signed-in to the Free plan, which gives you "
    "unlimited data on USA, Netherlands, and Japan servers.")
IMG_ADD("proton_a_01_signup.png",
        "Proton VPN sign-up screen on Android")

H3("Step 3: Allow VPN connection permission")
NUM("Tap Quick Connect or pick a country flag manually.")
NUM("Android will pop up a Connection request dialog the first time. "
    "Tap OK to allow Proton VPN to establish a VPN connection.")
NUM("A small key icon appears in the status bar — that means the VPN "
    "is active.")
IMG_ADD("proton_a_02_dlogin.png",
        "Login and download progress on the Android app")

H3("Step 4: Connect to a server in the US, UK or Netherlands")
NUM("Tap United States (or any flag from the free list).")
NUM("Wait for the green Connected text and the timer to start.")
NUM("Open a browser and visit whatismyipaddress.com — the country "
    "should now show as US/UK/Netherlands.")
IMG_ADD("proton_a_03_connect.png",
        "Connecting to a server on the Proton VPN Android app")

TIP("Always launch the VPN BEFORE opening CapCut, Play Store, or any "
    "geo-restricted app. Switching the VPN on after launching the "
    "target app often does not help — close and re-open the app.")

H3("Step 5: Tweak Android settings for stable VPN")
NUM("On Android 11+ go to Settings → Network and Internet → VPN.")
NUM("Tap the gear icon next to Proton VPN.")
NUM("Turn on Always-on VPN — this auto-reconnects if the connection drops.")
NUM("Optional: Turn on Block connections without VPN to prevent leaks.")

WARN("Block connections without VPN means apps will refuse to connect "
     "while the VPN is off. Toggle it back if you face issues with "
     "non-VPN apps.")

H2("1.3 Installing a VPN on Windows PC")

H3("Step 1: Download Proton VPN for Windows")
NUM("Open your browser and go to https://protonvpn.com/download.")
NUM("Click Get Proton VPN for Windows.")
NUM("Run the downloaded ProtonVPN_Setup.exe installer.")
NUM("Follow on-screen prompts. Default options work fine.")
IMG_ADD("vpn_05_download.png", "Proton VPN download prompt on the website")
IMG_ADD("vpn_06_install.png", "Windows installer for Proton VPN")

H3("Step 2: Sign in or create account")
NUM("Open Proton VPN from the Start menu.")
NUM("Sign in with the same account you created on mobile, or create a "
    "new free account.")
IMG_ADD("vpn_03_account.png", "Proton VPN account creation page")
IMG_ADD("vpn_04_password.png", "Generating a strong password")

H3("Step 3: Connect to a US/UK server")
NUM("In the app, click Quick Connect for an automatic best server, OR "
    "click the country list and pick United States or United Kingdom.")
NUM("Wait for the status to turn green.")
NUM("Now any app you open will see your computer as US/UK based.")

H2("1.4 Installing a VPN on iPhone / iPad")

H3("Quick Setup")
NUM("Open the App Store on your iPhone or iPad.")
NUM("Search Proton VPN.")
NUM("Tap Get → Install.")
NUM("Open the app, sign up or log in.")
NUM("Tap Quick Connect.")
NUM("iOS shows a permission dialog: Add VPN Configurations? — tap Allow.")
NUM("Enter your iPhone passcode / Face ID to confirm.")
NUM("Connection turns green. You can now open LumaFusion or any other "
    "app with the VPN active.")

NOTE("LumaFusion itself is NOT geo-restricted in India. You only need a "
     "VPN on iPhone/iPad if you want to access region-locked content "
     "stores or cloud services.")

# ============================================================
# PART 2 — CAPCUT
# ============================================================
H1("Part 2 — CapCut (PC + Android)")

H2("2.1 Installing CapCut in India")

H3("Method A — On a Windows PC")

P("Step 1: Change Windows region", bold=True)
NUM("Press Windows + I to open Settings.")
NUM("Go to Time & Language → Language & Region.")
NUM("Change Country or region from India to United Kingdom or United States.")

P("Step 2: Make sure your VPN is connected to UK/USA (Section 1.3).")

P("Step 3: Download CapCut", bold=True)
NUM("Open https://www.capcut.com in your browser.")
NUM("Click Download for Windows.")
IMG_ADD("capcut_01_site.png", "CapCut official website homepage")
IMG_ADD("capcut_02_download.png", "Download button for CapCut PC")
NUM("Run the .exe installer and follow prompts.")

P("Step 4: Sign in", bold=True)
NUM("Launch CapCut. Sign in with Google, TikTok or Facebook.")
IMG_ADD("capcut_03_signup.png", "CapCut account sign-up screen")

P("Step 5: Create your first project", bold=True)
IMG_ADD("capcut_04_create_project.png",
        "Click New Project on the CapCut home screen")
IMG_ADD("capcut_05_import.png",
        "Import clips by clicking Import or pressing Ctrl + I")
IMG_ADD("capcut_06_arrange.png",
        "Arrange the imported clips on the timeline")

WARN("Always download from capcut.com, NOT the Microsoft Store. The "
     "Pro export trick in Section 2.11 only works on the standalone "
     "installer.")

H3("Method B — On Android")

NUM("Connect Proton VPN to a UK or USA server (Section 1.2).")
NUM("Open Play Store → Settings → General → Account preferences → "
    "Country and profiles. Wait for the country to change.")
NUM("Clear Play Store cache (long-press Play Store icon → App info → "
    "Storage → Clear cache).")
NUM("Search CapCut on Play Store. Install it.")
NUM("Open CapCut with VPN still on, sign in.")
IMG_ADD("capcut_m_01_apk.png", "CapCut APK / Play Store search on mobile")
IMG_ADD("capcut_m_03_signup.png", "CapCut mobile sign-up")
IMG_ADD("capcut_m_04_project.png", "Create a new project on mobile")

YT("CapCut India install (Kripesh Adwani)",
   "https://kripeshadwani.com/how-to-use-capcut-in-india/")
YT("Install CapCut after India ban (YouTube)",
   "https://www.youtube.com/watch?v=NChj2zD9J-Q")

H2("2.2 CapCut interface tour")
P("CapCut PC has 4 main areas: top-left media library, top-right preview, "
  "bottom timeline, right-side properties panel.")
BUL("Ctrl + B — Split clip at playhead")
BUL("Ctrl + I — Import media")
BUL("Spacebar — Play / Pause")
BUL("Alt + K — Open keyframe panel")
BUL("Ctrl + Z / Y — Undo / Redo")

H2("2.3 Splitting and cutting clips")

H3("Step-by-step (with screenshots)")

NUM("Drag your video to the timeline. The clip appears as a long bar at "
    "the bottom.")
IMG_ADD("capcut_split_1.png",
        "Step 1 — Add the video clip to the CapCut timeline")

NUM("Drag the edge of the clip inwards if you want to trim from the start "
    "or end without splitting.")
IMG_ADD("capcut_split_2.png",
        "Step 2 — Drag the edge handles to trim the clip from start or end")

NUM("Move the playhead (the vertical white line) to the exact moment "
    "you want to cut.")
NUM("Click the Split icon (scissors) on the toolbar above the timeline. "
    "Or press Ctrl + B.")
IMG_ADD("capcut_split_4.png",
        "Step 3 — Position the playhead and press the scissors / Ctrl+B to split")

NUM("To delete only what is left or right of the playhead, use the "
    "Delete left and Delete right buttons in the toolbar.")
IMG_ADD("capcut_cut_delete.png",
        "Step 4 — Use Delete left / Delete right to remove segments instantly")

NUM("To remove a middle section: split at the start, split at the end, "
    "click the middle clip, press Delete on your keyboard. The remaining "
    "clips snap together.")

TIP("Cut on action and on the beat. If your video has music, split where "
    "the beat hits. Movement-based cuts feel professional even with no "
    "other effects.")

H3("On CapCut Mobile (Android / iOS)")
NUM("Tap the clip on the timeline.")
NUM("Tap Split in the bottom toolbar.")
NUM("Tap Delete to remove the unwanted section.")
IMG_ADD("capcut_cut_mobile.png",
        "Splitting on CapCut Mobile — Split and Delete options at the bottom")

YT("CapCut split clips beginner tutorial",
   "https://www.youtube.com/watch?v=eMt7IcYwe-k")
YT("Cut, trim, split (CapCut PC)",
   "https://www.youtube.com/watch?v=7Z8kCYKGTD0")
YT("Step-by-step trim and split (Minitool)",
   "https://moviemaker.minitool.com/news/trim-and-split-video-in-capcut.html")

H2("2.4 Keyframes for animation")

P("A keyframe is a marker that says: at this moment, this property has "
  "this value. Two keyframes with different values = animation. CapCut "
  "automatically generates the smooth motion in between.")

H3("Zoom-in animation (step-by-step)")
NUM("Click the clip on the timeline to select it. The right-side "
    "properties panel appears.")
NUM("Move the playhead to the start of the clip.")
NUM("In the right panel under Basic, look for the diamond icon next "
    "to Scale.")
NUM("Click the diamond — it turns orange. That is your first keyframe "
    "at 100% scale.")
NUM("Move the playhead 2 seconds forward by clicking on the timeline.")
NUM("Change the Scale value to 120%. CapCut automatically adds a second "
    "keyframe at this position.")
NUM("Press spacebar — the clip slowly zooms in from 100% to 120%.")

H3("Fade-in animation")
NUM("Move playhead to start of clip.")
NUM("Set Opacity = 0%, click the diamond.")
NUM("Move playhead 1 sec forward.")
NUM("Set Opacity = 100%. Done.")

H3("Other useful keyframes")
BUL("Position — text flying in from off-screen.")
BUL("Rotation — logo spinning 360 degrees as a stinger.")
BUL("Volume — duck music down when someone speaks.")

YT("CapCut keyframes step-by-step (PC)",
   "https://www.youtube.com/watch?v=nLER1f51FPw")
YT("Master keyframes like a pro",
   "https://www.youtube.com/watch?v=C563AeLl8eQ")

H2("2.5 Transitions between clips")

H3("Step-by-step")
NUM("Place two clips so they touch on the timeline.")
NUM("Click the Transitions tab in the top-left.")
NUM("Browse categories: Basic, Camera, Light Effect, Glitch, "
    "Distortion, Slide, Mask, MG, Social Media, Overlay, Blur.")
NUM("Hover to preview. Click the download icon if shown.")
NUM("Drag the transition between the two clips.")
NUM("Click the transition box → set Duration in the right panel.")

H3("Apply one transition to all clips")
NUM("Drag your favourite transition between the first two clips.")
NUM("Right-click the transition → Apply to all.")

NOTE("Maximum transition duration = half the length of the shorter "
     "neighbouring clip. A 4 sec + 3 sec pair allows a 1.5 sec max "
     "transition.")

YT("Apply same transition to all clips (PC)",
   "https://www.youtube.com/watch?v=LrwgdRaukJo")
YT("Add transitions in CapCut (CCT)",
   "https://contentcreatortemplates.com/learn/capcut/how-to-add-transitions-in-capcut")

H2("2.6 Effects and overlays")

H3("Add a video effect")
NUM("Click Effects → Video Effects.")
NUM("Browse: Trending, Basic, Bling, Vintage, Atmosphere, Fantasy.")
NUM("Click an effect to download and preview.")
NUM("Drag onto the timeline above your clip — it sits on its own track.")
NUM("Stretch or shrink the effect bar to control duration.")
NUM("Click the effect → adjust intensity, speed, color in the right panel.")

H3("Picture-in-Picture (PIP) overlay — step-by-step with screenshots")

NUM("Open your main video clip in CapCut.")
IMG_ADD("capcut_pip_1_import.png",
        "Step 1 — Import your main video clip into CapCut")

NUM("In the menu bar, look for the Overlay option (CapCut Mobile) or "
    "drag a second clip above the main clip on the timeline (CapCut PC).")
IMG_ADD("capcut_pip_2_open.png",
        "Step 2 — Open the Overlay feature from the menu / toolbar")

NUM("Tap or click + Add overlay and select the second video, image, "
    "or stock clip you want to layer on top.")
IMG_ADD("capcut_pip_3_add.png",
        "Step 3 — Pick the second clip to use as the overlay layer")

NUM("In the preview window, drag the overlay around to position it. "
    "Pinch the corners (mobile) or drag the corner handles (PC) to "
    "resize.")
IMG_ADD("capcut_pip_4_resize.png",
        "Step 4 — Reposition, resize, and rotate the PIP overlay on the canvas")

NUM("Customize the overlay further: in the right panel set Opacity, "
    "Blend mode (Screen for light leaks, Multiply for darks, Overlay "
    "for general), add a border, drop shadow, or animation.")
IMG_ADD("capcut_pip_5_edit.png",
        "Step 5 — Edit the overlay layer — opacity, blend mode, border, animation")

NUM("On the timeline, drag the ends of the overlay track to set when "
    "it appears and disappears.")

NUM("Export when satisfied (Section 13 has the export cheat sheet).")
IMG_ADD("capcut_pip_6_export.png",
        "Step 6 — Export the finished picture-in-picture video")

H3("Other types of overlay")
BUL("Stock light leaks / dust — set Blend mode to Screen.")
BUL("Rain or scratched film overlays — Blend mode Multiply.")
BUL("Logo / watermark — keep Opacity around 60-70%, no blend mode.")
BUL("Reaction-cam style PIP — small overlay in the corner, add a "
    "thin border and drop shadow.")

IMG_ADD("capcut_07_effects.png",
        "Effects and filters panel — apply animated effects on top of overlays")

YT("Add overlays in CapCut PC",
   "https://www.youtube.com/watch?v=DPGe5ziNrSo")
YT("Effects to overlay technique",
   "https://www.youtube.com/watch?v=BmZWXujR1Ik")
YT("Picture in Picture in CapCut PC — Lesson 12",
   "https://www.youtube.com/watch?v=xlBbqDtNwvM")
YT("PIP video effect on CapCut PC",
   "https://www.youtube.com/watch?v=9CABdQXf3pE")
YT("PIP in CapCut Desktop",
   "https://www.youtube.com/watch?v=jE4J4spI0JE")
YT("PIP on CapCut Mobile",
   "https://www.youtube.com/watch?v=O_yMNqt7cq8")

H2("2.7 Text masking")

H3("Video inside text (beginner)")
NUM("Drag video to timeline.")
NUM("Click Text → Add default text.")
NUM("Type a short word like BOLD. Use a thick font (Impact, Anton, Bebas).")
NUM("Make font huge, 200pt+.")
NUM("Stretch text bar to cover full clip.")
NUM("Apply Mask → use text as the mask source. Video plays inside the letters.")

H3("Walk-behind text reveal (advanced)")
NUM("Place clip on Track 1.")
NUM("Duplicate clip onto Track 3 above.")
NUM("Add text on Track 2 between the two video layers.")
NUM("Select Track 3 → Mask → Custom shape.")
NUM("Mask out only the part of the body that should cover the text.")
NUM("Use keyframes on the mask Position to follow the person frame by frame.")

TIP("Shoot with a tripod. Camera shake makes mask animation 10x harder.")

YT("Text masking (PC) — step-by-step",
   "https://www.youtube.com/watch?v=oJdVRVzLvVw")
YT("Reveal text mask tutorial",
   "https://www.youtube.com/watch?v=qDZoDwhEq5o")
YT("Mask text in CapCut (2025)",
   "https://www.youtube.com/watch?v=Hai5lrZDrkE")

H2("2.8 Color grading basics")
P("Click Adjust in the right panel to access these sliders. Work top to "
  "bottom:")
BUL("Brightness — pull up if dark, down if blown out.")
BUL("Contrast — most phone footage is flat; +10 to +20 starting point.")
BUL("Saturation — stay between -10 and +15.")
BUL("Sharpness — only push if soft; max +10.")
BUL("Highlights — pull down to recover sky details.")
BUL("Shadows — pull up to see detail without killing mood.")
BUL("Temperature — warm right for golden hour, cool left for moody.")
BUL("Tint — fix green / magenta color casts.")

YT("Color grading in CapCut (full tutorial)",
   "https://www.youtube.com/watch?v=q8tp6NGLBEM")

H2("2.9 HSL adjustment for cinematic color")
P("HSL = Hue, Saturation, Lightness. Lets you change one specific color "
  "without affecting the rest. The secret behind the teal-and-orange look.")

IMG_ADD("capcut_hsl_1.jpg", "Opening the HSL panel in CapCut")
IMG_ADD("capcut_hsl_3.jpg", "HSL color circles for individual hue control")

H3("Step-by-step")
NUM("Select your clip.")
NUM("Click Adjust → scroll to HSL.")
NUM("Click on a color circle (Red, Orange, Yellow, Green, Aqua, Blue, "
    "Purple, Pink).")
NUM("Three sliders appear: Hue (shift), Saturation (intensity), "
    "Lightness (brightness).")
NUM("Repeat for each color you want to control.")

IMG_ADD("capcut_hsl_4.jpg", "Adjusting Hue, Saturation, Lightness sliders")
IMG_ADD("capcut_hsl_5.jpg", "Targeted color shift with the HSL panel")

H3("Recipe: Teal & Orange cinematic look")
NUM("Click Orange. Saturation +15, Lightness +5 (pops skin tones).")
NUM("Click Yellow. Saturation -20, Hue toward Orange (kills sickly yellows).")
NUM("Click Aqua. Saturation +10.")
NUM("Click Blue. Saturation +20, Lightness -10 (deep skies and shadows).")
NUM("Click Green. Saturation -15, Hue toward Aqua.")
NUM("Toggle Adjust off and on to compare.")

IMG_ADD("capcut_hsl_8.jpg", "Before / after with HSL applied")

H3("Recipe: Moody filmic look")
NUM("Pull Highlights down, Shadows up slightly.")
NUM("Lower overall Saturation by 10-15.")
NUM("In HSL, lower Lightness on Reds and Oranges by -10.")
NUM("Push Blues toward Cyan, Lightness -10.")
NUM("Slight Tint to the left for green hint.")

TIP("Save your HSL recipe as a preset. The preset icon is in the Adjust "
    "panel — apply with one click on every clip.")

YT("HSL color grading explained",
   "https://filmora.wondershare.com/video-editor-review/hsl-capcut.html")

H2("2.10 Color gradations and pro looks")
P("Pro grading is layered:")
NUM("Step 1 Correction — fix exposure, white balance, contrast.")
NUM("Step 2 Base look — apply a Filter or LUT (Filters → Movie / "
    "Vlogger), reduce intensity to 40-60%.")
NUM("Step 3 HSL refinement — target individual colors as in 2.9.")
NUM("Step 4 Final polish — vignette, grain, +5 Sharpness.")

H3("Common color grades")
BUL("Warm Golden — outdoor, family, sunset. Temp +10, Highlights -5, Sat +10.")
BUL("Teal & Orange — talking head, travel. Use HSL recipe above.")
BUL("High Contrast B&W — emotional. Sat -100, Contrast +20.")
BUL("Pastel Soft — beauty, fashion Reels. Sat -15, Highlights -10.")
BUL("Faded Film — vintage. Use Faded preset, Shadows +20, low grain.")

H2("2.11 CapCut Pro export trick (free, no watermark)")
P("Source: https://youtu.be/Zm4i-YT5M1k (Tamil tutorial). The trick "
  "uses CapCut's pre-process engine to render Pro effects into a hidden "
  "cache folder, which you re-import.")

WARN("Workaround, not a hack. Personal learning use only. Buy CapCut Pro "
     "for commercial work.")

H3("Step 1 — Check version")
NUM("Open CapCut → click your profile → check version.")
NUM("Trick is most reliable on v5.x. Newer versions (7.x+) sometimes "
    "encrypt the cache.")

H3("Step 2 — Build edit with Pro effects")
NUM("Edit normally — cuts, transitions, text, music.")
NUM("Apply any Pro effects you want (they show a Pro tag).")
NUM("Preview to confirm.")

H3("Step 3 — Compound clip")
NUM("Press Ctrl + A to select all.")
NUM("Right-click → Create compound clip.")
NUM("Timeline shows one merged green block.")

H3("Step 4 — Apply Motion Blur and pre-process")
NUM("With the compound clip selected, go to the right-side properties "
    "panel for the video.")
NUM("Scroll down through the panel until you find Motion Blur.")
NUM("Slide the Motion Blur value all the way to 0.")
NUM("This forces CapCut to pre-process the compound clip with the Pro "
    "effects baked in.")
NUM("Wait for the progress bar to reach 100%. Do not move on until it "
    "is fully done.")

H3("Step 5 — Open the hidden cache folder")
NUM("Press Windows + R → type appdata → Enter.")
NUM("In the address bar, replace Roaming with Local. Press Enter.")
NUM("Navigate to: Local → CapCut → User Data → Projects → "
    "com.lveditor.draft → [Project Name] → Resources → Video Alg "
    "(this is the Video Algorithm folder).")
NUM("Inside Video Alg you will find the .mp4 file with all the Pro "
    "features rendered into it (ignore .alpha files).")

NOTE("If AppData is hidden in File Explorer: View → Show → Hidden items.")

H3("Step 6 — Re-import into CapCut and export")
NUM("Copy that .mp4 from the Video Alg folder to your Desktop or any "
    "easy-to-find location.")
NUM("Go back to CapCut.")
NUM("Press Ctrl + I and import the .mp4 you just copied.")
NUM("Drag the imported file onto the timeline (or onto the compound "
    "clip and choose Replace clip → Yes).")
NUM("Click Export in the top-right.")
NUM("Set High Quality — 1080p or 4K, MP4, 30 or 60 fps.")
NUM("Click Export. The Pro features are already baked into the file, "
    "so the export goes through with no watermark and all the Pro "
    "effects intact.")

H3("If it fails")
BUL("Make sure you used the standalone installer, not Microsoft Store.")
BUL("Try downgrading CapCut to v5.x.")
BUL("Make sure Motion Blur was set to 0 and pre-processing reached "
    "100% before grabbing the .mp4.")
BUL("Check the Video Alg folder, not Combination — newer CapCut "
    "versions store the rendered file under Video Alg.")
BUL("If the .mp4 plays without Pro effects, the cache was deleted or "
    "overwritten; redo Step 4 and grab the file immediately.")

YT("Original Tamil tutorial (Shaid's reference)",
   "https://youtu.be/Zm4i-YT5M1k")
YT("Pro effects free workaround (English)",
   "https://www.youtube.com/watch?v=UNBRq0sJ52U")
YT("CapCut Pro PC for free (no crack)",
   "https://www.youtube.com/watch?v=SYyIdhWdChs")
YT("Export Pro videos free",
   "https://www.youtube.com/watch?v=bIqA8SpkJfI")

# ============================================================
# PART 3 — CLIPCHAMP
# ============================================================
H1("Part 3 — Microsoft Clipchamp (Web + Windows 11)")

H2("3.1 What Clipchamp is and how it differs from CapCut")
P("Clipchamp is Microsoft's free video editor, pre-installed on "
  "Windows 11 and also available as a browser-based app at clipchamp.com. "
  "It is owned by Microsoft (acquired in 2021).")

P("Key differences from CapCut:", bold=True)
BUL("Browser-based — runs in Chrome / Edge with no install needed. "
    "CapCut PC is a desktop app.")
BUL("No keyframes for free users — Clipchamp has limited keyframe "
    "support. CapCut wins on animation control.")
BUL("Simpler interface — easier to learn, fewer pro features. Good for "
    "quick edits and slideshows.")
BUL("Stock library tied to Microsoft 365 — premium stock requires a "
    "subscription. Free users get a smaller pool.")
BUL("No region restrictions in India — Clipchamp works without VPN.")
BUL("Color grading is limited to brightness, contrast, saturation, "
    "exposure, temperature. NO HSL panel and NO masks.")

P("When to use Clipchamp:", bold=True)
BUL("Quick edits where you do not need keyframes or HSL.")
BUL("Family videos, slideshows, simple corporate clips.")
BUL("If your laptop is low on disk space (it is browser-based).")
BUL("If you are inside the Microsoft 365 ecosystem already.")

P("When NOT to use it:", bold=True)
BUL("Cinematic color grading (use CapCut).")
BUL("Complex motion graphics with keyframes (use CapCut or LumaFusion).")
BUL("Mobile-first editing on phone (use CapCut mobile or LumaFusion).")

H2("3.2 Getting started")
NUM("Open Edge or Chrome, go to https://clipchamp.com.")
NUM("Click Try for free — sign in with Microsoft, Google or email.")
NUM("Click Create a new video.")
NUM("Choose aspect ratio: 16:9 (YouTube), 9:16 (Reels), 1:1 (Insta feed).")

IMG_ADD("clip_01_overview.png",
        "The Clipchamp web editor interface")

H2("3.3 Splitting clips")

P("Clipchamp lays out the editor in 4 areas: left toolbar (media, "
  "templates, text, transitions, effects), top floating toolbar, right "
  "property panel, bottom timeline.")

IMG_ADD("clip_02_uidiagram.png",
        "Labelled diagram of Clipchamp's UI sections")
IMG_ADD("clip_03_lefttool.png",
        "Left toolbar — media, templates, text, transitions, effects")
IMG_ADD("clip_04_rightpanel.png",
        "Right property panel — captions, audio, filters, color, speed")
IMG_ADD("clip_07_timeline.png",
        "Timeline at the bottom — cuts, playback, zoom controls")

H3("Step-by-step split — with screenshots")

NUM("Click Import media (top-left) and add your video file.")
NUM("Drag the video onto the timeline at the bottom.")
NUM("Click the clip on the timeline to select it (it gets a blue border).")

NUM("Move the seeker (the vertical playhead line) to the exact moment "
    "you want to cut.")
IMG_ADD("clip_split_1_seeker.png",
        "Step 1 — Move the seeker on the timeline to your cut point")

NUM("Click the Split icon (scissors) on the floating toolbar above the "
    "timeline. Keyboard shortcut: press S.")
IMG_ADD("clip_split_2_scissors.png",
        "Step 2 — Click the scissors icon (or press S) to split the clip")

NUM("The clip is now broken into two parts on the timeline.")
IMG_ADD("clip_split_3_two_parts.png",
        "Step 3 — Clip is now split into two separate clips")

NUM("To remove a section: right-click the unwanted clip and choose "
    "Delete (or click and press the Delete key).")
IMG_ADD("clip_split_4_delete.png",
        "Step 4 — Delete the unwanted segment to clean up the timeline")

NUM("Repeat as many times as needed. Drag clips left/right to "
    "rearrange them.")

YT("Clipchamp split and trim (Kevin Stratvert)",
   "https://www.youtube.com/watch?v=pHMoPh6mzm4")
YT("Microsoft Support — split videos and audio",
   "https://support.microsoft.com/en-us/topic/how-to-split-or-cut-videos-and-audio-assets-fe2d10f0-1ebc-4b8b-858e-4f9442356425")
YT("Clipchamp tutorial — cut and split",
   "https://windowsreport.com/clipchamp-how-to-cut-video/")

H2("3.4 Keyframes (and the workaround)")

WARN("Important: Clipchamp DOES NOT support full keyframe animation like "
     "CapCut. You cannot keyframe Scale, Position, Opacity, Rotation, or "
     "Volume manually. This is the single biggest limitation vs CapCut.")

H3("Workarounds for animation in Clipchamp")
BUL("Use built-in Animate options on text and images (Fade in, Slide in, "
    "Zoom in, Bounce). Click your asset → Animate tab → choose preset.")
BUL("For video zooms: split your clip into 2 parts, apply different "
    "Crop / Zoom levels via Filters, and put a transition between them.")
BUL("Use Speed ramping (right panel → Speed) to fake dynamic moves.")
BUL("If you really need keyframes, edit the asset in CapCut and re-import "
    "the result into Clipchamp.")

P("Steps to apply built-in animation:", bold=True)
NUM("Click any text or image asset on the timeline.")
NUM("Click Animate in the right panel (or floating toolbar).")
NUM("Pick from In, Out, or Both presets.")
NUM("Adjust Duration with the slider.")

H2("3.5 Transitions")

H3("Step-by-step")
NUM("Place two clips next to each other on the timeline.")
NUM("Click Transitions in the left toolbar.")
NUM("Categories: Crossfade, Fade to Black, Wipe, Spin, Slide, Heart, "
    "Zoom, etc.")
NUM("Drag the chosen transition between the two clips.")

IMG_ADD("clip_t_04_drag.png", "Dragging a transition between two clips")
IMG_ADD("clip_t_05_apply.png", "Transition applied at the cut point")

H3("Adjust duration")
NUM("Click the transition box on the timeline.")
NUM("In the right panel, drag the Duration slider (0.1s to 2s typical).")

IMG_ADD("clip_t_06_duration.png", "Adjusting transition duration in the property panel")

YT("Clipchamp transitions tutorial",
   "https://www.youtube.com/watch?v=rUsoT976QyE")
YT("Add new transitions",
   "https://www.youtube.com/watch?v=aRYf94TNzGc")

H2("3.6 Effects and filters")

H3("Apply a filter")
NUM("Click your clip on the timeline.")
NUM("In the right panel, click Filters tab.")
NUM("Hover to preview, click to apply: Vintage, B&W, Sepia, Vignette, etc.")
NUM("Drag the Intensity slider.")

H3("Add an effect")
NUM("Effects in Clipchamp are simpler than CapCut — Blur, Slow Zoom, "
    "Old Film, Mirror, Glitch, etc.")
NUM("Click Effects tab in the right panel.")
NUM("Click an effect to apply.")

NOTE("Clipchamp does not have body effects or animated overlay effects "
     "like CapCut. For that, edit on CapCut.")

H2("3.7 Text and titles")

NUM("Click Text in the left toolbar.")
NUM("Browse styles: Plain, Animated, Subtitle, Lower third.")
NUM("Click a style to add it. It lands on the timeline above your video.")
NUM("Double-click in the preview to type your text.")
NUM("In the right panel, change Font, Size, Color, Position, Animation.")
NUM("Stretch the text bar on timeline to control how long it shows.")

YT("Clipchamp text and titles",
   "https://www.youtube.com/watch?v=AnfIJda3E1M")

H2("3.8 Color adjustments")

H3("Step-by-step")
NUM("Click your clip.")
NUM("In the right panel, click Adjust colors.")
NUM("Sliders available:")
BUL("Exposure")
BUL("Contrast")
BUL("Saturation")
BUL("Temperature")
BUL("Vibrance")
NUM("Adjust each slider. Preview updates live.")

NOTE("There is NO HSL panel, NO color wheels, NO masks. For cinematic "
     "color grading you need CapCut or LumaFusion.")

H3("Filters as a shortcut")
P("Use a filter for a quick base look (Cinema, Sunny, Cool, Vivid), "
  "then adjust the sliders to fine-tune. This is the closest you can "
  "get to a graded look in Clipchamp.")

H2("3.9 Exporting")

NUM("Click Export in the top-right.")
NUM("Choose resolution: 480p, 720p, or 1080p (free).")
NUM("4K export requires Clipchamp Premium.")
NUM("File downloads as MP4 to your computer.")

IMG_ADD("clip_t_07_export.png", "Clipchamp export dialog")

NOTE("Clipchamp has no watermark in the free version — you do not need "
     "a workaround like CapCut. This is a real strength.")

# ============================================================
# PART 4 — LUMAFUSION
# ============================================================
H1("Part 4 — LumaFusion (iPhone / iPad / Android)")

H2("4.1 Why LumaFusion is different")

P("LumaFusion is a paid, professional-grade video editor for iPhone, "
  "iPad and (newer) Android. It is the closest thing on mobile to "
  "Final Cut Pro or Premiere.")

P("Key differences:", bold=True)
BUL("Paid app — one-time purchase around USD 30 (~₹2,500). No subscription.")
BUL("Up to 6 video tracks + 6 audio tracks (CapCut mobile is 2-3 typically).")
BUL("Full keyframe control on every parameter.")
BUL("Built-in color wheels, scopes (vectorscope, waveform), LUT support.")
BUL("J-cuts and L-cuts — the standard pro audio-video offset technique.")
BUL("Magic motion tracking, chroma key, masks.")
BUL("No watermark, no Pro tier, no ads. Pay once, own it.")

P("When to use LumaFusion:", bold=True)
BUL("Filmmaking-style edits on iPad.")
BUL("Multi-cam interviews and podcasts.")
BUL("Travel videos with serious color grading.")
BUL("If you want to leave CapCut behind for something more pro.")

H2("4.2 Getting LumaFusion")

H3("On iPhone / iPad")
NUM("Open the App Store.")
NUM("Search LumaFusion.")
NUM("Tap Buy — it's a paid app.")
NUM("Once purchased, download and open.")
NUM("No VPN needed. Available in India.")

H3("On Android")
NUM("Open Play Store.")
NUM("Search LumaFusion.")
NUM("Tap Install (paid).")
NUM("Note: Android version is newer and missing a few iOS features.")

H3("On Chromebook")
NUM("Available via the Play Store or Chromebook download badge on "
    "luma-touch.com.")

NOTE("There is a separate add-on called Luma Fx and Storyblocks integration "
     "for an extra fee — both optional.")

H2("4.3 Interface and timeline")

IMG_ADD("luma_01_main.png",
        "The LumaFusion main editing screen on iPad",
        width_inches=6.0)

P("LumaFusion's screen layout:")
BUL("Top-left: Source preview window.")
BUL("Top-right: Program preview (your edit).")
BUL("Center: Project Library with media bins.")
BUL("Bottom: Timeline with up to 6 video tracks + 6 audio tracks.")
BUL("Right edge: Toolbar (cut, mark, add transitions, sync, share).")

IMG_ADD("luma_cb_03.png",
        "LumaFusion timeline with trim and edit tools",
        width_inches=6.0)

H2("4.4 Splitting and trimming")

P("LumaFusion's interface is iPad-first, so the split workflow uses "
  "your finger on the timeline.")

IMG_ADD("luma_tutsplus_1_interface.png",
        "LumaFusion main interface — Source preview, Program preview, and Timeline",
        width_inches=5.5)

H3("Step-by-step")
NUM("In the Library (centre panel), find your imported clip.")
IMG_ADD("luma_tutsplus_2_files.png",
        "Browse and pick a clip from the LumaFusion Library / Files panel")

NUM("Drag the clip from the Library onto the timeline.")
NUM("Tap the clip to select it — yellow handles appear at both ends.")
NUM("Drag the timeline left or right with your finger to position the "
    "blue playhead at the cut point.")
NUM("Tap the Scissors / Split icon in the timeline toolbar (right edge).")
NUM("The clip is split into two clips at that moment.")

P("Quick split shortcut:", bold=True)
P("Tap two fingers on either side of the blue playhead — LumaFusion "
  "creates a cut at the playhead position.", italic=True)

H3("Trim a clip without splitting")
NUM("Tap the clip to select it.")
NUM("Drag the yellow handle on the left edge inwards to trim the start.")
NUM("Drag the yellow handle on the right edge inwards to trim the end.")
NUM("Or pinch the clip with two fingers — LumaFusion zooms / trims based "
    "on your gesture.")

H3("Delete a middle section")
NUM("Split at the start of the section.")
NUM("Split at the end.")
NUM("Tap the middle clip to select it.")
NUM("Tap the trash icon in the toolbar.")
NUM("LumaFusion automatically closes the gap (Ripple Delete).")

H3("J-cuts and L-cuts (audio offset)")
P("This is the technique that makes interviews feel cinematic — audio "
  "from clip B starts before clip B's video appears (J-cut), or video "
  "cuts away while clip A's audio continues (L-cut).")
NUM("Right-click (long press) the clip → Detach audio.")
NUM("Audio is now on its own track.")
NUM("Trim the video early or late while leaving audio in place.")

YT("LumaFusion timeline navigation",
   "https://luma-touch.com/tutorials/")
YT("J-cuts and L-cuts in LumaFusion",
   "https://luma-touch.com/tutorials/")

H2("4.5 Keyframes")

P("LumaFusion has keyframe support on every clip parameter. Tap the "
  "diamond icon next to any parameter to add a keyframe.")

IMG_ADD("luma_03_keyframe.png",
        "Keyframe panel in LumaFusion — every property is animatable",
        width_inches=6.0)

H3("Step-by-step zoom-in keyframe")
NUM("Tap the clip on the timeline.")
NUM("Tap the Frame and Fit icon (square with arrows).")
NUM("Move the playhead to the start.")
NUM("Tap the Add keyframe (diamond) icon.")
NUM("Move the playhead 2 seconds forward.")
NUM("Pinch zoom in the preview to 1.2x scale.")
NUM("Keyframe is added automatically.")
NUM("Tap Done. Play to confirm the zoom animation.")

TIP("LumaFusion auto-adds a keyframe whenever you change a value while "
    "Add keyframe is enabled — so it's faster than CapCut once you "
    "are used to it.")

YT("LumaFusion keyframe basics",
   "https://luma-touch.com/tutorials/")

H2("4.6 Transitions")

H3("Step-by-step")
NUM("Place two clips touching on the timeline.")
NUM("Tap the cut point between them.")
NUM("Tap Add transition in the toolbar.")
NUM("Pick: Cross-dissolve, Fade to Black, Slide, Wipe, Push, Mosaic, etc.")
NUM("Drag the corner of the transition to adjust duration on the timeline.")

H3("Apply same transition to many cuts")
NUM("Long-press the transition.")
NUM("Tap Apply to all → confirm.")

YT("10 easy transitions in LumaFusion",
   "https://www.youtube.com/watch?v=em87yXYEDnc")

H2("4.7 Effects")

H3("Built-in effects")
NUM("Tap the clip → Effects tab.")
NUM("Categories: Color, Stylize, Distort, Blur, Sharpen, Mosaic, Chroma key.")
NUM("Tap an effect → adjust sliders → tap Done.")
NUM("Effects can be keyframed — animate the intensity over time.")

H3("Chroma key (green screen)")
NUM("Place green screen clip on Track 2.")
NUM("Tap the clip → Effects → Chroma Key.")
NUM("Tap the eyedropper, tap the green area in the preview.")
NUM("Adjust Strength, Softness, Spill suppression.")
NUM("Done — green is removed.")

H2("4.8 Color grading and LUTs")

P("LumaFusion has the most professional color tools on mobile.")

H3("Color preset")
NUM("Tap the clip → Color & Effects.")
NUM("Tap Color Preset → pick from Cinematic, Vintage, B&W, etc.")
NUM("Adjust intensity slider.")

H3("Manual color grading")
NUM("Tap Color & Effects → Color Adjustments.")
NUM("Sliders: Exposure, Contrast, Highlights, Shadows, Whites, Blacks, "
    "Saturation, Temperature, Tint, Sharpness.")
NUM("Below that you have Color Wheels: Lift (shadows), Gamma (midtones), "
    "Gain (highlights). Drag the puck to push color into that range.")

H3("Apply a LUT (.cube file)")
NUM("Tap Color & Effects → Add Effect → LUT.")
NUM("Tap Choose LUT → import a .cube file from Files app or iCloud.")
NUM("Adjust intensity 0-100%.")

H3("Scopes")
P("Tap the gear icon in the preview → enable Scopes. You get vectorscope, "
  "waveform, RGB parade — pro-level color reading tools that CapCut "
  "does not have.")

YT("LumaFusion color grading 101",
   "https://academy.iphoneographers.tv/p/lumafusion-color-grading-101")
YT("LUTs in LumaFusion",
   "https://filmora.wondershare.com/video-editing/lumafusion-luts.html")

NOTE("LumaFusion does not have a single HSL panel like CapCut, but the "
     "combination of Color Wheels + Scopes + LUTs gives you more "
     "precise control overall.")

H2("4.9 Audio levels and metering")

P("LumaFusion has built-in audio meters — something casual editors "
  "ignore until their video sounds bad. Always glance at the meters "
  "before exporting.")

IMG_ADD("luma_tutsplus_3_meters.png",
        "Audio meters in LumaFusion — voice should peak around -6dB",
        width_inches=5.5)

P("Targets:", bold=True)
BUL("Voice peaks: around -6 dB (yellow zone, never red).")
BUL("Music bed: -18 to -20 dB so it does not overpower voice.")
BUL("Master output: never higher than -3 dB to avoid clipping.")

H2("4.10 Exporting")

NUM("Tap the Share icon (top-right toolbar).")
NUM("Choose Movie.")
NUM("Settings:")
BUL("Resolution: 720p, 1080p, 4K.")
BUL("Frame rate: 24, 25, 30, 50, 60 fps.")
BUL("Quality: Low, Medium, High, Best.")
BUL("Format: H.264 or HEVC.")
NUM("Tap Save to Photos, Files, or share to YouTube / Vimeo / Frame.io.")

# ============================================================
# PART 5 — COMPARISON
# ============================================================
H1("Part 5 — Tool Comparison Cheat Sheet")

table = doc.add_table(rows=1, cols=4)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Feature"
hdr[1].text = "CapCut"
hdr[2].text = "Clipchamp"
hdr[3].text = "LumaFusion"

rows = [
    ("Cost", "Free + Pro tier", "Free (Premium for 4K)", "Paid one-time (~$30)"),
    ("Platforms", "PC, Mac, Android, iOS", "Web, Win 11", "iOS, iPad, Android, Chromebook"),
    ("Available in India", "VPN needed", "Yes (no VPN)", "Yes (no VPN)"),
    ("Watermark on free", "Yes (use trick)", "No", "N/A"),
    ("Keyframes", "Yes (full)", "Limited (Animate presets)", "Yes (full, every parameter)"),
    ("Transitions", "100+ animated", "30+ basic", "20+ pro"),
    ("Text masking", "Yes", "No", "Yes (via masks)"),
    ("HSL color", "Yes", "No", "No (uses wheels + scopes)"),
    ("Color wheels", "Yes", "No", "Yes"),
    ("LUT support", "Yes (Pro tier)", "No", "Yes"),
    ("Multi-track", "Up to 10", "Limited", "6 video + 6 audio"),
    ("Best for", "Reels, TikTok, social", "Quick edits, slideshows", "Pro mobile filmmaking"),
    ("Learning curve", "Medium", "Low", "Medium-High"),
]
for r in rows:
    cells = table.add_row().cells
    for i, v in enumerate(r):
        cells[i].text = v

doc.add_paragraph()
P("Recommendation:", bold=True)
BUL("Start on CapCut — it covers 80% of what most creators need.")
BUL("Use Clipchamp on Windows 11 for fast, no-watermark cuts of "
    "presentations and family content.")
BUL("Move to LumaFusion if you are serious about iPad-based "
    "filmmaking and want pro color tools.")

# ============================================================
# PART 6 — YOUTUBE LIBRARY
# ============================================================
H1("Part 6 — YouTube Tutorial Library")

H2("CapCut")
YT("CapCut for PC complete tutorial (Primal Video)",
   "https://www.youtube.com/watch?v=pHMoPh6mzm4")
YT("CapCut beginners — split, trim, basics",
   "https://www.youtube.com/watch?v=eMt7IcYwe-k")
YT("Keyframes step-by-step (PC)",
   "https://www.youtube.com/watch?v=nLER1f51FPw")
YT("Keyframes — text and clip animation 2025",
   "https://www.youtube.com/watch?v=YKJc2YXmmKc")
YT("Master keyframes like a pro",
   "https://www.youtube.com/watch?v=C563AeLl8eQ")
YT("Apply same transition to all clips",
   "https://www.youtube.com/watch?v=LrwgdRaukJo")
YT("Color grading in CapCut",
   "https://www.youtube.com/watch?v=q8tp6NGLBEM")
YT("Text masking on CapCut PC",
   "https://www.youtube.com/watch?v=oJdVRVzLvVw")
YT("Reveal text mask",
   "https://www.youtube.com/watch?v=qDZoDwhEq5o")
YT("Add overlays in CapCut PC",
   "https://www.youtube.com/watch?v=DPGe5ziNrSo")
YT("Effects to overlay",
   "https://www.youtube.com/watch?v=BmZWXujR1Ik")
YT("CapCut Pro for free PC (Tamil) — Shaid's reference",
   "https://youtu.be/Zm4i-YT5M1k")
YT("Pro effects free workaround (English)",
   "https://www.youtube.com/watch?v=UNBRq0sJ52U")
YT("Export Pro videos free",
   "https://www.youtube.com/watch?v=bIqA8SpkJfI")

H2("Clipchamp")
YT("Clipchamp tutorial 2025 (Primal Video)",
   "https://primalvideo.com/guides/microsoft-clipchamp-tutorial-for-beginners-the-complete-guide/")
YT("Clipchamp transitions tutorial",
   "https://www.youtube.com/watch?v=rUsoT976QyE")
YT("Clipchamp transition effects",
   "https://www.youtube.com/watch?v=AnfIJda3E1M")
YT("New transitions in Clipchamp",
   "https://www.youtube.com/watch?v=aRYf94TNzGc")
YT("Best Clipchamp tips and tricks (Kevin Stratvert)",
   "https://kevinstratvert.com/2023/01/12/best-clipchamp-video-editing-tips-and-tricks/")

H2("LumaFusion")
YT("LumaFusion tutorial — Primal Video",
   "https://primalvideo.com/video-creation/editing/lumafusion-tutorial-how-to-edit-video-on-iphone-ipad/")
YT("Official LumaFusion tutorials hub",
   "https://luma-touch.com/tutorials/")
YT("10 easy transitions and effects",
   "https://www.youtube.com/watch?v=em87yXYEDnc")
YT("LumaFusion color grading 101",
   "https://academy.iphoneographers.tv/p/lumafusion-color-grading-101")
YT("LumaFusion essentials masterclass",
   "https://www.skillshare.com/en/classes/lumafusion-essentials-masterclass-2025-3-5hrs-new/928031644")

H2("VPN setup")
YT("Use CapCut in India (Kripesh Adwani)",
   "https://kripeshadwani.com/how-to-use-capcut-in-india/")
YT("Best VPN for CapCut (Comparitech)",
   "https://www.comparitech.com/blog/vpn-privacy/best-vpns-for-capcut/")
YT("Best free VPN for Android (TechRadar)",
   "https://www.techradar.com/vpn/best-free-vpn")

# ============================================================
# PART 7 — COMMON MISTAKES
# ============================================================
H1("Part 7 — Common Mistakes to Avoid")
BUL("Skipping the VPN before first launch in India. CapCut PC will "
    "throw a Network Error and refuse to load.")
BUL("Using the Microsoft Store version of CapCut and then trying the "
    "Pro export trick. The trick only works on the standalone .exe.")
BUL("Over-using transitions. One or two styles per video — max.")
BUL("Pushing Saturation and Sharpness too far. Subtle wins.")
BUL("Forgetting to check audio levels before exporting. Voice peaks "
    "around -6 dB, music around -18 dB.")
BUL("Using thin fonts for text masks. Always go thick (Impact, Anton, "
    "Bebas, Anton).")
BUL("Forgetting to deselect a keyframe before changing a property "
    "globally — you may create accidental keyframes.")
BUL("Editing on a phone screen and not checking on a laptop. Colors "
    "look different on bigger screens.")
BUL("Trying to use Clipchamp for keyframe-heavy edits. It does not "
    "have full keyframes — switch to CapCut.")
BUL("Ignoring scopes in LumaFusion when grading. They tell you when "
    "you are clipping highlights or crushing blacks.")

# ============================================================
# PART 8 — REFERENCES
# ============================================================
H1("Part 8 — References & Sources")

P("Reference video used for the CapCut Pro export trick:")
BUL("https://youtu.be/Zm4i-YT5M1k — How to use CapCut Pro for free in PC (Tamil)")

P("CapCut")
BUL("https://www.capcut.com/resource/how-to-use-capcut-on-pc — Official PC guide")
BUL("https://www.capcut.com/resource/how-to-add-keyframes-in-capcut — Keyframes")
BUL("https://www.capcut.com/tools/hsl-color — HSL adjustment")
BUL("https://www.capcut.com/resource/capcut-color-grading — Color grading")
BUL("https://www.tourboxtech.com/en/news/capcut-transitions.html — Transitions")
BUL("https://kripeshadwani.com/how-to-use-capcut-in-india/ — India VPN install")
BUL("https://techydruid.com/get-capcut-pro-free-pc/ — Pro free trick (PC)")
BUL("https://www.ruhanirabin.com/unlock-capcut-pro-windows-free/ — Pro free trick variations")
BUL("https://filmora.wondershare.com/video-editor-review/hsl-capcut.html — HSL deep dive")
BUL("https://www.alphr.com/capcut-how-to-use-keyframes/ — AlphR keyframes guide")
BUL("https://www.createthat.ai/blog/color-grading-in-capcut — Color panels overview")

P("Clipchamp")
BUL("https://clipchamp.com/en/blog/how-to-navigate-video-editing-tools/ — Interface overview")
BUL("https://clipchamp.com/en/blog/add-transitions-videos-slideshows/ — Transitions guide")
BUL("https://primalvideo.com/guides/microsoft-clipchamp-tutorial-for-beginners-the-complete-guide/ — Beginner guide")
BUL("https://www.capcut.com/resource/clipchamp-vs-capcut — Comparison")

P("LumaFusion")
BUL("https://luma-touch.com/tutorials/ — Official tutorials")
BUL("https://www.podfeet.com/blog/2021/01/lumafusion/ — Detailed review with screenshots")
BUL("https://www.creativebloq.com/how-to/lumafusion-ipad-basics — iPad basics")
BUL("https://academy.iphoneographers.tv/p/lumafusion-color-grading-101 — Color grading course")
BUL("https://primalvideo.com/video-creation/editing/lumafusion-tutorial-how-to-edit-video-on-iphone-ipad/ — Primal Video tutorial")

P("VPN")
BUL("https://protonvpn.com/free-vpn/android — Proton VPN Android (official)")
BUL("https://www.techradar.com/vpn/best-free-vpn — Best free VPN 2026")
BUL("https://www.pcworld.com/article/695190/best-free-vpn-for-android.html — Best Android VPN")
BUL("https://thebestvpn.com/best-vpns-android/ — Top 5 Android VPNs tested")

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run("End of guide. Verify, then this exports to PDF.")
fr.italic = True
fr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Size: {os.path.getsize(OUT)} bytes")
